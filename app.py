import os
import pandas as pd
from flask import Flask, render_template, request, send_file, jsonify
from docxtpl import DocxTemplate
import uuid
from datetime import datetime

# 创建 Flask 应用
app = Flask(__name__)

# 配置文件上传
# 使用 /tmp 目录（Render 的临时存储）
UPLOAD_FOLDER = '/tmp/uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key')

def generate_proposal(data_df):
    """
    根据数据生成方案书
    返回生成的 Word 文件路径
    """
    # 使用绝对路径
    template_path = 'proposal_template.docx'
    
    # 如果模板不存在，创建一个基础模板
    if not os.path.exists(template_path):
        from docx import Document
        doc = Document()
        doc.add_heading('项目方案书', 0)
        doc.add_paragraph('项目名称：{{ project_name }}')
        doc.add_paragraph('负责人：{{ manager }}')
        doc.add_paragraph('预算金额：{{ budget }} 万元')
        doc.add_paragraph('项目周期：{{ duration }} 个月')
        doc.add_heading('项目概述', level=1)
        doc.add_paragraph('{{ overview }}')
        doc.add_heading('项目清单', level=1)
        doc.add_paragraph('{% for item in items %}')
        doc.add_paragraph('- {{ item.name }}：{{ item.desc }}')
        doc.add_paragraph('{% endfor %}')
        doc.save(template_path)
    
    # 处理数据：如果有多行，生成汇总信息
    context = {
        'project_name': data_df.iloc[0]['项目名称'] if '项目名称' in data_df.columns else '未命名项目',
        'manager': data_df.iloc[0]['负责人'] if '负责人' in data_df.columns else '待定',
        'budget': data_df['预算'].sum() if '预算' in data_df.columns else 0,
        'duration': data_df['周期'].mean() if '周期' in data_df.columns else 0,
        'overview': '本项目旨在提升业务效率，通过系统化建设实现数字化转型。',
        'items': []
    }
    
    # 生成项目清单
    for _, row in data_df.iterrows():
        context['items'].append({
            'name': row.get('项目名称', '未命名'),
            'desc': row.get('描述', '无详细说明')
        })
    
    # 生成唯一文件名
    output_filename = f"proposal_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    
    # 渲染模板
    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_path)
    
    return output_path

@app.route('/')
def index():
    """上传页面"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """处理上传文件并生成方案书"""
    try:
        # 检查是否有文件上传
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        # 检查文件格式
        if not file.filename.endswith(('.xlsx', '.xls', '.csv')):
            return jsonify({'error': '只支持 Excel 或 CSV 文件'}), 400
        
        # 保存上传的文件
        input_filename = f"input_{uuid.uuid4().hex[:8]}{os.path.splitext(file.filename)[1]}"
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
        file.save(input_path)
        
        # 读取数据表
        if input_path.endswith('.csv'):
            df = pd.read_csv(input_path)
        else:
            df = pd.read_excel(input_path)
        
        # 检查数据是否为空
        if df.empty:
            return jsonify({'error': '数据表为空'}), 400
        
        # 生成方案书
        output_path = generate_proposal(df)
        
        # 返回下载链接
        return jsonify({
            'success': True,
            'download_url': f'/download/{os.path.basename(output_path)}',
            'message': '方案书生成成功！'
        })
        
    except Exception as e:
        return jsonify({'error': f'生成失败：{str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """下载生成的方案书"""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': '文件不存在'}), 404
    
    # 自定义下载文件名
    download_name = f"方案书_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# 定时清理临时文件（可选）
@app.after_request
def cleanup(response):
    # 可以添加清理逻辑，比如删除超过1小时的临时文件
    return response

if __name__ == '__main__':
    # 开发环境
    app.run(debug=True, host='0.0.0.0', port=5000)